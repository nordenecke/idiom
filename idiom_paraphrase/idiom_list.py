# -*- coding: utf-8 -*-
"""
Created on Sat Jul 15 21:58:21 2017

@author: norden
"""

import os

from docx import Document
from docx.shared import Pt
#from docx.shared import Inches
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
#from docx.oxml.shared import OxmlElement, qn

import winreg
def get_desktop():
    key = winreg.OpenKey(winreg.HKEY_CURRENT_USER,\
                          r'Software\Microsoft\Windows\CurrentVersion\Explorer\Shell Folders',)
    return winreg.QueryValueEx(key, "Desktop")[0]


idioms_list=[]
c_lst=[]

output_column_number=3
output_row_number=5



def get_idiom_list(filename):
    desktop_path=get_desktop()
    with open(desktop_path+"\\"+filename,"r",encoding='utf-8') as idioms_src:
        for line in idioms_src:
            idioms_list.append(line.strip())
#            print(line.strip())
#        print(idioms_list)
        return idioms_list

def put_content(filename, content_list):
    with open(filename,"w+") as content_dst:
        content_dst.truncate()
        for i in range(len(content_list)):
            content_dst.write(content_list[i].idiom+"\n")
            content_dst.write(content_list[i].phonetic_symbol+"\n")
            content_dst.write(content_list[i].paraphrase+"\n")
            content_dst.write("***************************\n")
        return True
    
def add_idiom_and_phonetic_symbol_paragraph(cell,idiom,phonetic_symbol):
    #add idiom
    paragraph1=cell.add_paragraph(u"")
    paragraph1.alignment=WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph1.add_run(idiom)
    if len(idiom) <= 7:
        run.font.size = Pt(48)
    elif len(idiom) <= 9:
        run.font.size = Pt(36)
    elif len(idiom) <= 12:
        run.font.size = Pt(28)
    elif len(idiom) <= 14:
        run.font.size = Pt(24)
    elif len(idiom) <= 15:
        run.font.size = Pt(22)
    elif len(idiom) <= 16:
        run.font.size = Pt(20)
    elif len(idiom) <= 18:
        run.font.size = Pt(18)
    else:
        run.font.size = Pt(16)
        
    run.font.name = 'Consolas'
    run.bold=True


    #add phonetic_symbol
    paragraph2=cell.add_paragraph(u"")
    paragraph2.alignment=WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph2.add_run(phonetic_symbol)
    run.font.size = Pt(16)
    run.font.name = 'Consolas'
    run.bold=True

    
def add_paraphrase_paragraph(cell,paraphrase):
    #add idiom
    paragraph1=cell.add_paragraph(u"")
    paragraph1.alignment=WD_ALIGN_PARAGRAPH.LEFT
    #set font size
    run = paragraph1.add_run(paraphrase)
    run.font.size = Pt(12)
    run.font.name=u'微软雅黑'
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')

    #set font
#    run = paragraph1.add_run(u'Set Font.')
#    run.font.name = 'Consolas'

    #set chinese font
#    run = paragraph1.add_run(u'set chinese font')
#    run.font.name=u'宋体'
#    r = run._element
#    r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

    #set italic
#    run = paragraph1.add_run(u'set italic、')
#    run.italic = True

    #set bold
#    run = paragraph1.add_run(u'bold').bold = False
    
def set_raws_height(rows):
    for i in range(len(rows)):
        tr=rows[i]._tr
        trPr=tr.get_or_add_trPr()
        trHeight = OxmlElement('w:trHeight')
        trHeight.set(qn('w:val'), "3032")
        trHeight.set(qn('w:hRule'), "atLeast")
        trPr.append(trHeight)  



def set_cell_vertical_alignment(cell, align="center"): 
    try:   
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcValign = OxmlElement('w:vAlign')  
        tcValign.set(qn('w:val'), align)  
        tcPr.append(tcValign)
        return True 
    except:
        traceback.print_exc()             
        return False
    
def put_docx(filename,content_list):
    if len(content_list)==0:
        return False

    #open document
    document = Document()
    #加入不同等级的标题
    document.add_heading(u'Idioms Paraphrase',0)
#    document.add_heading(u'一级标题',1)
#    document.add_heading(u'二级标题',2)
    #添加文本
    paragraph = document.add_paragraph(u'成语释义')
    #设置字号
#    run = paragraph.add_run(u'设置字号、')
#    run.font.size = Pt(24)

    #设置字体
#    run = paragraph.add_run('Set Font,')
#    run.font.name = 'Consolas'

    #设置中文字体
#    run = paragraph.add_run(u'设置中文字体、')
#    run.font.name=u'宋体'
#    r = run._element
#    r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
#    run.font.size = Pt(24)    

    #设置斜体
#    run = paragraph.add_run(u'斜体、')
#    run.italic = True

    #设置粗体
#    run = paragraph.add_run(u'粗体').bold = True

    #增加引用
#    document.add_paragraph('Intense quote', style='Intense Quote')

    #增加无序列表
#    document.add_paragraph(
#        u'无序列表元素1', style='List Bullet'
#    )
#    document.add_paragraph(
#        u'无序列表元素2', style='List Bullet'
#    )
    #增加有序列表
#    document.add_paragraph(
#        u'有序列表元素1', style='List Number'
#    )
#    document.add_paragraph(
#        u'有序列表元素2', style='List Number'
#    )
    #增加图像（此处用到图像image.bmp，请自行添加脚本所在目录中）
#    document.add_picture('image.bmp', width=Inches(1.25))
    for i in range(len(content_list)):
        for j in range(len(content_list[i])):
            if j==0:
                run = paragraph.add_run(u'\n\n'+content_list[i][j])
                run.font.name=u'宋体'
                r = run._element
                r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
                run.font.size = Pt(24)
#            else:
            elif j==2:
                run = paragraph.add_run(u'\n'+content_list[i][j])
                run.font.name=u'宋体'
                r = run._element
                r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
                run.font.size = Pt(11)

    #保存文件
    desktop_path=get_desktop()
    if os.path.exists(desktop_path+"\\"+filename):
        os.remove(desktop_path+"\\"+filename)
#    Sections = document.sections
#    for section in Sections:
#        section.top_margin = 914400/8
#        section.bottom_margin = 914400/8
#        section.left_margin = 1143000/8
#        section.right_margin = 1143000/8
    document.save(desktop_path+"\\"+filename)


def main():
    print("This is get_idiom_list main function!")
#    print(get_idiom_list("idioms.txt"))
#    c_lst.append(idiom_content_structure.idiom_content("111","bbb","ccc"))
#    c_lst.append(idiom_content_structure.idiom_content("222","bbb","ccc"))
#    put_content("idioms_content.txt",c_lst)


if __name__=="__main__":
    main()
    
    